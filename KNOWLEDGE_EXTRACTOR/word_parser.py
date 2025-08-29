import os
import sys
from pathlib import Path
from typing import Optional, List, Dict, Any
import logging

# Word document processing libraries
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("python-docx not found. Install with: pip install python-docx")

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

class WordExtractor:
    """Advanced Word document text extractor with error handling."""
    
    def __init__(self, docx_path: str):
        self.docx_path = Path(docx_path)
        self.document = None
        self.text_content = {}
        
    def validate_file(self) -> bool:
        """Validate Word document file exists and is accessible."""
        if not self.docx_path.exists():
            logger.error(f"Word document not found: {self.docx_path}")
            return False
        
        if not self.docx_path.is_file():
            logger.error(f"Path is not a file: {self.docx_path}")
            return False
            
        if self.docx_path.stat().st_size == 0:
            logger.error(f"Word document is empty: {self.docx_path}")
            return False
            
        # Check if it's a .docx file
        if self.docx_path.suffix.lower() not in ['.docx', '.doc']:
            logger.warning(f"File may not be a Word document: {self.docx_path}")
            
        return True
    
    def load_document(self) -> bool:
        """Load Word document with error handling."""
        try:
            self.document = Document(self.docx_path)
            logger.info(f"Word document loaded successfully. Paragraphs: {len(self.document.paragraphs)}")
            return True
            
        except Exception as e:
            logger.error(f"Failed to load Word document: {e}")
            return False
    
    def extract_text_from_paragraphs(self) -> str:
        """Extract text from all paragraphs."""
        text = ""
        
        try:
            for paragraph in self.document.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
                    
            logger.info(f"Extracted text from {len(self.document.paragraphs)} paragraphs")
            return text.strip()
            
        except Exception as e:
            logger.error(f"Failed to extract text from paragraphs: {e}")
            return ""
    
    def extract_text_from_tables(self) -> str:
        """Extract text from all tables."""
        text = ""
        
        try:
            for table in self.document.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text += " | ".join(row_text) + "\n"
                text += "\n"  # Add space between tables
                
            logger.info(f"Extracted text from {len(self.document.tables)} tables")
            return text.strip()
            
        except Exception as e:
            logger.error(f"Failed to extract text from tables: {e}")
            return ""
    
    def extract_document_properties(self) -> Dict[str, str]:
        """Extract document properties/metadata."""
        properties = {
            "title": "",
            "author": "",
            "subject": "",
            "keywords": "",
            "comments": "",
            "category": "",
            "created": "",
            "modified": ""
        }
        
        try:
            core_props = self.document.core_properties
            
            if core_props.title:
                properties["title"] = core_props.title
            if core_props.author:
                properties["author"] = core_props.author
            if core_props.subject:
                properties["subject"] = core_props.subject
            if core_props.keywords:
                properties["keywords"] = core_props.keywords
            if core_props.comments:
                properties["comments"] = core_props.comments
            if core_props.category:
                properties["category"] = core_props.category
            if core_props.created:
                properties["created"] = str(core_props.created)
            if core_props.modified:
                properties["modified"] = str(core_props.modified)
                
        except Exception as e:
            logger.warning(f"Failed to extract document properties: {e}")
            
        return properties
    
    def extract_all_text(self) -> Dict[str, Any]:
        """Extract all text from Word document with comprehensive metadata."""
        if not self.validate_file():
            return {"error": "Invalid Word document file"}
        
        if not self.load_document():
            return {"error": "Failed to load Word document"}
        
        # Extract text from different sources
        paragraph_text = self.extract_text_from_paragraphs()
        table_text = self.extract_text_from_tables()
        
        # Combine all text
        full_text = ""
        if paragraph_text:
            full_text += paragraph_text + "\n\n"
        if table_text:
            full_text += "--- TABLES ---\n" + table_text + "\n\n"
        
        full_text = full_text.strip()
        
        result = {
            "file_path": str(self.docx_path),
            "total_paragraphs": len(self.document.paragraphs),
            "total_tables": len(self.document.tables),
            "paragraphs": {},
            "tables": {},
            "full_text": full_text,
            "metadata": self.extract_document_properties()
        }
        
        # Extract individual paragraphs with formatting info
        for i, paragraph in enumerate(self.document.paragraphs):
            result["paragraphs"][i + 1] = {
                "text": paragraph.text,
                "style": paragraph.style.name if paragraph.style else "Normal",
                "has_text": bool(paragraph.text.strip()),
                "runs": len(paragraph.runs)
            }
        
        # Extract individual tables
        for i, table in enumerate(self.document.tables):
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text.strip())
                table_data.append(row_data)
            
            result["tables"][i + 1] = {
                "rows": len(table.rows),
                "columns": len(table.columns) if table.rows else 0,
                "data": table_data
            }
        
        return result
    
    def save_extracted_text(self, output_path: Optional[str] = None) -> str:
        """Save extracted text to a file."""
        result = self.extract_all_text()
        
        if "error" in result:
            logger.error(f"Cannot save: {result['error']}")
            return ""
        
        if not output_path:
            output_path = self.docx_path.with_suffix('.txt')
        
        try:
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(f"Word Document Text Extraction Results\n")
                f.write(f"File: {result['file_path']}\n")
                f.write(f"Paragraphs: {result['total_paragraphs']}\n")
                f.write(f"Tables: {result['total_tables']}\n")
                f.write("=" * 50 + "\n\n")
                f.write(result['full_text'])
            
            logger.info(f"Text saved to: {output_path}")
            return str(output_path)
            
        except Exception as e:
            logger.error(f"Failed to save text: {e}")
            return ""

def extract_word_text(file_path: str) -> Dict[str, Any]:
    """
    Extract text from a single Word document file.
    
    Args:
        file_path: Path to the Word document file
        
    Returns:
        Dict containing extraction results with keys:
        - success: Boolean indicating if extraction was successful
        - file_path: Original file path
        - text: Extracted text content
        - metadata: Document metadata if available
        - paragraphs: Paragraph-by-paragraph extraction details
        - tables: Table extraction details
        - error: Error message if extraction failed
    """
    try:
        extractor = WordExtractor(file_path)
        result = extractor.extract_all_text()
        
        if "error" in result:
            return {
                "success": False,
                "file_path": file_path,
                "error": result["error"]
            }
        
        return {
            "success": True,
            "file_path": file_path,
            "text": result["full_text"],
            "metadata": result["metadata"],
            "paragraphs": result["paragraphs"],
            "tables": result["tables"],
            "total_paragraphs": result["total_paragraphs"],
            "total_tables": result["total_tables"]
        }
        
    except Exception as e:
        logger.error(f"Failed to extract text from {file_path}: {e}")
        return {
            "success": False,
            "file_path": file_path,
            "error": str(e)
        }

def process_batch_word_docs(file_paths: List[str]) -> List[Dict[str, Any]]:
    """
    Process multiple Word document files in batch.
    
    Args:
        file_paths: List of file paths to process
        
    Returns:
        List of extraction results for each file
    """
    results = []
    total_files = len(file_paths)
    
    logger.info(f"Starting batch processing of {total_files} Word documents")
    
    for i, file_path in enumerate(file_paths, 1):
        logger.info(f"Processing file {i}/{total_files}: {file_path}")
        result = extract_word_text(file_path)
        results.append(result)
        
        if result["success"]:
            logger.info(f"✓ Successfully processed: {file_path}")
        else:
            logger.warning(f"✗ Failed to process: {file_path} - {result['error']}")
    
    # Summary statistics
    successful = sum(1 for r in results if r["success"])
    failed = total_files - successful
    
    logger.info(f"Batch processing complete: {successful} successful, {failed} failed")
    
    return results

def extract_resume_sections(text: str) -> Dict[str, str]:
    """
    Extract structured sections from resume text.
    
    Args:
        text: Raw resume text
        
    Returns:
        Dict with structured sections (skills, experience, education, etc.)
    """
    sections = {
        "contact_info": "",
        "skills": "",
        "experience": "",
        "education": "",
        "summary": "",
        "other": ""
    }
    
    # Simple section extraction using keywords
    lines = text.split('\n')
    current_section = "other"
    
    for line in lines:
        line_lower = line.lower().strip()
        
        # Detect sections based on keywords
        if any(keyword in line_lower for keyword in ['skill', 'technology', 'programming', 'framework']):
            current_section = "skills"
        elif any(keyword in line_lower for keyword in ['experience', 'work', 'employment', 'job']):
            current_section = "experience"
        elif any(keyword in line_lower for keyword in ['education', 'degree', 'university', 'college', 'school']):
            current_section = "education"
        elif any(keyword in line_lower for keyword in ['summary', 'profile', 'objective', 'about']):
            current_section = "summary"
        elif any(keyword in line_lower for keyword in ['email', 'phone', '@', 'linkedin', 'github']):
            current_section = "contact_info"
        
        # Add line to current section
        if line.strip():
            sections[current_section] += line + "\n"
    
    # Clean up sections
    for key in sections:
        sections[key] = sections[key].strip()
    
    return sections

def main():
    """Main function for command line usage."""
    if len(sys.argv) > 1:
        docx_path = sys.argv[1]
        result = extract_word_text(docx_path)
        
        if result["success"]:
            print(f"✓ Successfully extracted text from: {docx_path}")
            print(f"Text length: {len(result['text'])} characters")
            print(f"Paragraphs: {result['total_paragraphs']}")
            print(f"Tables: {result['total_tables']}")
        else:
            print(f"✗ Failed to extract text: {result['error']}")
    else:
        print("Usage: python word_parser.py <file_path>")
        print("For batch processing, use the programmatic functions directly.")

if __name__ == "__main__":
    main()
